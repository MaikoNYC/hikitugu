"""Document file generation service (PDF/Word export)."""

import io
import re


class FileGeneratorService:
    """Generates PDF and Word files from document sections."""

    async def generate_pdf(self, document: dict, sections: list[dict]) -> bytes:
        """Generate a PDF file from document sections.

        Args:
            document: Document metadata dict.
            sections: List of section dicts with title and content.

        Returns:
            PDF file bytes.
        """
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)

        # Title page
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 24)
        pdf.cell(0, 60, "", ln=True)
        title = document.get("title", "Document")
        pdf.cell(0, 20, self._sanitize_text(title), ln=True, align="C")
        pdf.set_font("Helvetica", "", 12)
        target = document.get("target_user_email", "")
        if target:
            pdf.cell(0, 10, f"Target: {target}", ln=True, align="C")
        created = document.get("created_at", "")
        if created:
            pdf.cell(0, 10, f"Created: {created[:10]}", ln=True, align="C")

        # Section pages
        for section in sections:
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 16)
            section_title = section.get("title", "")
            pdf.cell(0, 12, self._sanitize_text(section_title), ln=True)
            pdf.ln(4)

            pdf.set_font("Helvetica", "", 11)
            content = section.get("content", "") or ""
            plain = self._markdown_to_plain(content)
            for line in plain.split("\n"):
                pdf.multi_cell(0, 6, self._sanitize_text(line))

        output = io.BytesIO()
        pdf.output(output)
        return output.getvalue()

    async def generate_docx(self, document: dict, sections: list[dict]) -> bytes:
        """Generate a Word (.docx) file from document sections.

        Args:
            document: Document metadata dict.
            sections: List of section dicts with title and content.

        Returns:
            DOCX file bytes.
        """
        from docx import Document as DocxDocument

        doc = DocxDocument()
        title = document.get("title", "Document")
        doc.add_heading(title, level=0)

        target = document.get("target_user_email", "")
        if target:
            doc.add_paragraph(f"対象者: {target}")
        created = document.get("created_at", "")
        if created:
            doc.add_paragraph(f"作成日: {created[:10]}")

        doc.add_page_break()

        for section in sections:
            section_title = section.get("title", "")
            doc.add_heading(section_title, level=1)
            content = section.get("content", "") or ""
            for line in content.split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=2)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=3)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif re.match(r"^\d+\.\s", stripped):
                    text = re.sub(r"^\d+\.\s", "", stripped)
                    doc.add_paragraph(text, style="List Number")
                else:
                    cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
                    cleaned = re.sub(r"\*(.+?)\*", r"\1", cleaned)
                    doc.add_paragraph(cleaned)

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    @staticmethod
    def _sanitize_text(text: str) -> str:
        """Replace characters unsupported by basic PDF fonts with ASCII equivalents."""
        return text.encode("latin-1", errors="replace").decode("latin-1")

    @staticmethod
    def _markdown_to_plain(md: str) -> str:
        """Strip basic markdown formatting to plain text."""
        text = re.sub(r"#{1,6}\s*", "", md)
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
        return text
